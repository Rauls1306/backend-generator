# metodologia_prisma.py
from typing import Dict

import openai
from docx import Document


def _gpt(prompt: str, max_tokens: int = 1200, temperature: float = 0.3) -> str:
    """
    Wrapper simple para ChatCompletion (openai==0.28).
    Asume que openai.api_key ya está configurada en main.py.
    """
    resp = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    return resp.choices[0].message.content.strip()


def _build_prisma_summary_text(prisma_stats: Dict[str, Dict[str, int]]) -> str:
    """
    prisma_stats: {
      "Scopus": {"identificados": 30, "incluidos": 9, "excluidos": 21},
      "SciELO": {...},
      ...
    }
    Devuelve un texto compacto para incluir en el prompt.
    """
    lines = []
    for base, stats in prisma_stats.items():
        ident = stats.get("identificados", 0)
        incl = stats.get("incluidos", 0)
        excl = stats.get("excluidos", 0)
        lines.append(
            f"- {base}: {ident} registros identificados, {excl} excluidos, {incl} artículos incluidos en la revisión."
        )
    return "\n".join(lines)


def generar_metodologia_y_figura_prisma(
    ruta_articulo: str,
    tema: str,
    pais: str,
    indexacion: str,
    prisma_stats: Dict[str, Dict[str, int]],
) -> Dict[str, str]:
    """
    1) Genera texto de Metodología adaptado al estudio usando un modelo estándar.
    2) Genera texto descriptivo para la Figura 1 (diagrama PRISMA).
    3) Inserta la Metodología en el Word del artículo (al final) con el título 'Metodología'.

    prisma_stats: {
      "Scopus": {"identificados": 30, "incluidos": 9, "excluidos": 21},
      "SciELO": {...},
      ...
    }

    Retorna:
    {
      "metodologia": <texto_metodologia>,
      "figura_prisma": <texto_figura>
    }
    """
    prisma_resumen_texto = _build_prisma_summary_text(prisma_stats)

    prompt = f"""
Eres un experto en redacción científica en español. A partir de los datos del estudio y del resumen del flujo PRISMA, debes redactar:

1. La sección 'Metodología' de un artículo científico de revisión bibliográfica.
2. Un texto descriptivo para la 'Figura 1. Diagrama de flujo PRISMA del proceso de selección de artículos'.

DATOS DEL ESTUDIO
- Tipo de trabajo: revisión bibliográfica.
- Tema central: {tema}.
- País de enfoque principal: {pais}.
- Nivel de indexación: {indexacion}.
- Resumen numérico PRISMA por base de datos:
{prisma_resumen_texto}

INSTRUCCIONES PARA LA METODOLOGÍA
- El texto debe seguir el estilo de un modelo estándar de metodología en revisiones bibliográficas:
  - Mencionar que se trata de una revisión bibliográfica de tipo {indexacion} (según corresponda).
  - Describir brevemente las bases de datos empleadas (por ejemplo: Scopus/ScienceDirect, SciELO, Dialnet).
  - Indicar el rango temporal aproximado de búsqueda (por ejemplo: últimos 7 años o 2018–2025).
  - Explicar que se aplicaron criterios de inclusión y exclusión (año, tipo de estudio, temática, idioma y acceso al texto completo).
  - Mencionar que se utilizó el diagrama PRISMA para resumir el proceso de identificación, cribado, exclusión e inclusión de artículos.
  - Conectar el número total de artículos incluidos con los datos del flujo PRISMA (sin listar uno por uno).

- Redacta la metodología en 3 a 5 párrafos, en tiempo pasado, con lenguaje académico claro y coherente.
- No inventes nombres de autores ni títulos específicos; limítate a describir el proceso metodológico.

INSTRUCCIONES PARA LA FIGURA PRISMA
- Redacta un texto breve (1 párrafo) que describa el contenido de la Figura 1, mencionando:
  - Que se trata de un diagrama de flujo PRISMA.
  - Que muestra los registros identificados, los excluidos y los artículos finales incluidos en la revisión.
  - Integra los números globales a partir del resumen PRISMA proporcionado (suma total de identificados y de incluidos).

FORMATO DE RESPUESTA (MUY IMPORTANTE)
Devuelve el resultado en este formato EXACTO:

METODOLOGIA:
[escribe aquí el texto completo de la sección metodología]

FIGURA_PRISMA:
[escribe aquí el texto descriptivo de la Figura 1]

No agregues ningún otro encabezado ni comentarios fuera de este formato.
"""

    respuesta = _gpt(prompt)

    # Parsear respuesta
    metodologia_texto = ""
    figura_texto = ""

    # Buscamos los marcadores METODOLOGIA: y FIGURA_PRISMA:
    parts = respuesta.split("FIGURA_PRISMA:")
    if len(parts) == 2:
        parte_metodo = parts[0]
        parte_fig = parts[1]

        # Quitar "METODOLOGIA:" del inicio de parte_metodo
        if "METODOLOGIA:" in parte_metodo:
            metodologia_texto = parte_metodo.split("METODOLOGIA:", 1)[1].strip()
        else:
            metodologia_texto = parte_metodo.strip()

        figura_texto = parte_fig.strip()
    else:
        # Si el modelo no respetó el formato (fallback)
        metodologia_texto = respuesta.strip()
        figura_texto = ""

    # Insertar en el DOCX del artículo
    try:
        doc = Document(ruta_articulo)

        # Título "Metodología"
        doc.add_heading("METODOLOGIA", level=1)

        for parrafo in metodologia_texto.split("\n"):
            if parrafo.strip():
                doc.add_paragraph(parrafo.strip())

        # Texto de la figura (como pie de figura / descripción)
        if figura_texto:
            doc.add_paragraph("")
            doc.add_paragraph("Figura 1. Diagrama de flujo PRISMA del proceso de selección de artículos.", style=None)
            doc.add_paragraph(figura_texto)

        doc.save(ruta_articulo)
    except Exception as e:
        print(f"❌ Error insertando metodología en el DOCX: {e}")

    return {
        "metodologia": metodologia_texto,
        "figura_prisma": figura_texto,
    }
