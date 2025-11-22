# prisma_processor.py
import os
from datetime import datetime
from typing import Dict, List

import openai
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader


def gpt(prompt: str, max_tokens: int = 1200, temperature: float = 0.3) -> str:
    """Wrapper simple para ChatCompletion (openai==0.28)."""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    return response.choices[0].message.content.strip()


def extract_text_from_pdf(pdf_path: str, max_chars: int = 16000) -> str:
    """
    Intenta extraer texto como PDF.
    Si falla (EOF marker, PDF corrupto, en realidad HTML, etc.),
    lee el archivo como texto plano (HTML) y devuelve los primeros max_chars.
    """
    text = ""

    # 1) Intento normal como PDF
    try:
        reader = PdfReader(pdf_path)
        pages_text = []
        for page in reader.pages:
            try:
                t = page.extract_text() or ""
            except Exception:
                t = ""
            pages_text.append(t)
        text = "\n".join(pages_text)
    except Exception as e:
        print(f"⚠️ No se pudo leer como PDF ({pdf_path}): {e}. Usando fallback a texto/HTML.")

    # 2) Si no se pudo sacar texto como PDF o quedó vacío, fallback a texto/HTML
    if not text.strip():
        try:
            with open(pdf_path, "rb") as f:
                raw = f.read()
            # Intentar decodificar como UTF-8, ignorando errores
            text = raw.decode("utf-8", errors="ignore")
        except Exception as e2:
            print(f"❌ Error leyendo archivo como texto plano ({pdf_path}): {e2}")
            return ""

    # 3) Recortar para no pasarle demasiadas cosas a GPT
    if len(text) > max_chars:
        return text[:max_chars]
    return text


PRISMA_PROMPT_BASE = """A partir de este artículo, extrae la siguiente información y preséntala en texto plano:

Objetivo del artículo: redáctalo en exactamente 30 palabras.
Metodología del artículo: descríbela en exactamente 30 palabras.
Tipo de metodología: elige solo entre las siguientes opciones: Cualitativa, Cuantitativa, Revisión sistemática o Mixta.
Resultados/aportes: descríbelos en prosa, con exactamente 60 palabras. Enfócate en los aportes centrales del estudio sin incluir interpretaciones ajenas.
Conclusiones: resume en exactamente 50 palabras.
País de origen: identifica un país específico. Si hay varios países o el estudio es global, escribir “Internacional”. Nunca dejar vacío este campo.
Año de publicación: debe ser un año específico. Si el artículo presenta un rango de años, elegir el más reciente.
Referencia en formato APA con DOI: redactar la referencia completa en estilo APA 7, incluyendo el DOI.
Cita APA abreviada: aplicar estas reglas:
- Un solo apellido por autor.
- Un autor: (Apellido, año).
- Dos autores: (Apellido y Apellido, año).
- Tres o más autores: (Apellido et al., año).
- Si los autores tienen varios apellidos, usar solo el primero.
Título completo del artículo.

ENTREGA LA INFORMACIÓN EN EL SIGUIENTE FORMATO EXACTO:

Artículo {idx}
Objetivo:
[texto]
Metodología:
[texto]
Tipo de metodología:
[texto]
Resultados/aportes:
[texto]
Conclusiones:
[texto]
País:
[texto]
Año:
[texto]
Referencia APA:
[texto]
Cita APA:
[texto]
Título:
[texto]

No agregues explicaciones adicionales, ni comentarios, ni texto fuera de este esquema.
"""


def build_prisma_prompt(article_text: str, idx: int) -> str:
    return (
        PRISMA_PROMPT_BASE.format(idx=idx)
        + "\n\nTEXTO DEL ARTÍCULO (recortado):\n\n"
        + article_text
    )


def parse_prisma_block(raw: str) -> Dict[str, str]:
    """Parsea el bloque devuelto por GPT al diccionario esperado."""
    lines = [l.strip() for l in raw.splitlines() if l.strip()]
    data: Dict[str, str] = {
        "objetivo": "",
        "metodologia": "",
        "tipo_metodologia": "",
        "resultados": "",
        "conclusiones": "",
        "pais": "",
        "anio": "",
        "referencia_apa": "",
        "cita_apa": "",
        "titulo": "",
    }
    label_map = {
        "objetivo": "objetivo",
        "metodología": "metodologia",
        "metodologia": "metodologia",
        "tipo de metodología": "tipo_metodologia",
        "tipo de metodologia": "tipo_metodologia",
        "resultados/aportes": "resultados",
        "resultados": "resultados",
        "conclusiones": "conclusiones",
        "país": "pais",
        "pais": "pais",
        "año": "anio",
        "ano": "anio",
        "referencia apa": "referencia_apa",
        "cita apa": "cita_apa",
        "título": "titulo",
        "titulo": "titulo",
    }

    current_key = None
    for line in lines:
        lower = line.lower()
        if lower.startswith("artículo") or lower.startswith("articulo"):
            continue

        matched_label = None
        for label in label_map.keys():
            prefix = label + ":"
            if lower.startswith(prefix):
                matched_label = label
                content = line[len(prefix):].strip()
                key = label_map[label]
                data[key] = content
                current_key = key
                break

        if matched_label is None:
            # línea de continuación del campo actual
            if current_key:
                if data[current_key]:
                    data[current_key] += " " + line
                else:
                    data[current_key] = line

    return data


def process_pdf_for_prisma(pdf_path: str, idx: int) -> Dict[str, str]:
    text = extract_text_from_pdf(pdf_path)
    prompt = build_prisma_prompt(text, idx)
    raw = gpt(prompt)
    data = parse_prisma_block(raw)
    return data


def procesar_pdfs_por_fuente(pdf_paths: List[str]) -> List[Dict[str, str]]:
    articulos: List[Dict[str, str]] = []
    for i, path in enumerate(pdf_paths, start=1):
        try:
            result = process_pdf_for_prisma(path, i)
            result["pdf_path"] = path
            articulos.append(result)
        except Exception as e:
            print(f"❌ Error procesando {path}: {e}")
    return articulos


def generar_word_prisma(
    indexacion: str,
    fuente: str,
    articulos: List[Dict[str, str]],
    output_dir: str = "output",
) -> str:
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    encabezado = f"PRISMA – {indexacion}"
    if fuente:
        encabezado += f" – {fuente}"
    doc.add_heading(encabezado, level=0)

    for i, art in enumerate(articulos, start=1):
        doc.add_paragraph(f"Artículo {i}")
        doc.add_paragraph(f"Objetivo:\n{art.get('objetivo','')}")
        doc.add_paragraph(f"Metodología:\n{art.get('metodologia','')}")
        doc.add_paragraph(f"Tipo de metodología:\n{art.get('tipo_metodologia','')}")
        doc.add_paragraph(f"Resultados/aportes:\n{art.get('resultados','')}")
        doc.add_paragraph(f"Conclusiones:\n{art.get('conclusiones','')}")
        doc.add_paragraph(f"País:\n{art.get('pais','')}")
        doc.add_paragraph(f"Año:\n{art.get('anio','')}")
        doc.add_paragraph(f"Referencia APA:\n{art.get('referencia_apa','')}")
        doc.add_paragraph(f"Cita APA:\n{art.get('cita_apa','')}")
        doc.add_paragraph(f"Título:\n{art.get('titulo','')}")
        doc.add_paragraph("")  # espacio

    filename = (
        f"PRISMA_{indexacion.replace(' ', '_')}_"
        f"{fuente.replace(' ', '_') if fuente else 'general'}_"
        f"{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    )
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def generar_excel_prisma(
    indexacion: str,
    fuente: str,
    articulos: List[Dict[str, str]],
    output_dir: str = "output",
) -> str:
    os.makedirs(output_dir, exist_ok=True)
    rows = []
    for art in articulos:
        rows.append(
            {
                "Objetivo": art.get("objetivo", ""),
                "Metodología": art.get("metodologia", ""),
                "Tipo de metodología": art.get("tipo_metodologia", ""),
                "Resultados/aportes": art.get("resultados", ""),
                "Conclusiones": art.get("conclusiones", ""),
                "País": art.get("pais", ""),
                "Año": art.get("anio", ""),
                "Referencia APA": art.get("referencia_apa", ""),
                "Cita APA": art.get("cita_apa", ""),
                "Título": art.get("titulo", ""),
            }
        )
    df = pd.DataFrame(rows)
    filename = (
        f"PRISMA_{indexacion.replace(' ', '_')}_"
        f"{fuente.replace(' ', '_') if fuente else 'general'}_"
        f"{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
    )
    filepath = os.path.join(output_dir, filename)
    df.to_excel(filepath, index=False)
    return filepath
