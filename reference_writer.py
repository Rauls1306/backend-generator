# reference_writer.py
import os
import re
from datetime import datetime

import openai
from docx import Document


def resumir_titulo(titulo: str, max_palabras: int = 10) -> str:
    palabras = titulo.split()
    if len(palabras) <= max_palabras:
        return titulo
    return " ".join(palabras[:max_palabras])


def gpt(prompt: str, max_tokens: int = 1200, temperature: float = 0.3) -> str:
    """
    Wrapper simple para generar texto con GPT-4 (openai==0.28).
    La API key ya debe estar configurada en main.py.
    """
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
            max_tokens=max_tokens,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("❌ Error en GPT (reference_writer):", str(e))
        raise


def generar_refs_institucionales(titulo: str, nivel: str, pais: str) -> dict:
    """
    Genera referencias INSTITUCIONALES para:
      - mundial
      - latam
      - nacional
    simulando lo descrito en la Segunda Parte:
    - Informes de organismos oficiales
    - Años 2020–2025
    - Formato APA 7
    """
    resultados = {}

    # 1) Problemática mundial
    prompt_mundial = (
        "Genera UNA lista de 4 referencias en formato APA 7, escritas en español, "
        "sobre la problemática global del siguiente tema de revisión bibliográfica:\n\n"
        f"Título: {titulo}\n\n"
        "Condiciones:\n"
        "- Las referencias deben ser INFORMES o DOCUMENTOS INSTITUCIONALES de organismos internacionales "
        "(por ejemplo: OMS/WHO, UNESCO, CEPAL, Banco Mundial, OCDE, UNICEF, ONU, etc.).\n"
        "- Años de publicación entre 2020 y 2025.\n"
        "- Incluye, cuando sea posible, la URL o DOI al final de cada referencia.\n"
        "- No agregues explicaciones adicionales, solo la lista de referencias, una por línea."
    )
    refs_mundial_raw = gpt(prompt_mundial)
    resultados["mundial"] = [r.strip() for r in refs_mundial_raw.split("\n") if r.strip()]

    # 2) Problemática latinoamericana
    prompt_latam = (
        "Genera UNA lista de 4 referencias en formato APA 7, escritas en español, "
        "sobre la problemática de este tema en América Latina:\n\n"
        f"Título: {titulo}\n\n"
        "Condiciones:\n"
        "- Referencias de organismos regionales o internacionales que aborden específicamente América Latina "
        "(ej.: CEPAL, OPS, BID, UNESCO, Banco Mundial, etc.).\n"
        "- Años de publicación entre 2020 y 2025.\n"
        "- Formato APA 7, con URL o DOI cuando sea posible.\n"
        "- Sin comentarios ni texto extra, solo la lista."
    )
    refs_latam_raw = gpt(prompt_latam)
    resultados["latam"] = [r.strip() for r in refs_latam_raw.split("\n") if r.strip()]

    # 3) Problemática nacional
    prompt_nacional = (
        "Genera UNA lista de 4 referencias en formato APA 7, escritas en español, "
        "sobre la problemática de este tema en el contexto nacional del país indicado.\n\n"
        f"Título del tema: {titulo}\n"
        f"País: {pais}\n\n"
        "Condiciones:\n"
        "- Referencias de ministerios, institutos nacionales de estadística, superintendencias, "
        "u otros organismos oficiales del país (por ejemplo: Ministerio de Salud, Ministerio de Educación, "
        "Instituto Nacional de Estadística, Superintendencias, etc.).\n"
        "- Años de publicación entre 2020 y 2025.\n"
        "- Formato APA 7, con URL cuando sea posible.\n"
        "- Solo la lista, una referencia por línea."
    )
    refs_nacional_raw = gpt(prompt_nacional)
    resultados["nacional"] = [r.strip() for r in refs_nacional_raw.split("\n") if r.strip()]

    return resultados


def generar_refs_teoricas_y_conceptuales(variable1: str, variable2: str) -> dict:
    """
    Genera referencias ACADÉMICAS (artículos científicos) para:
      - Teoría 1 (asociada a variable1)
      - Teoría 2 (asociada a variable2)
      - Variable 1 (concepto)
      - Variable 2 (concepto)
    Siguiendo la lógica de la Segunda Parte:
      - 3 referencias para cada teoría
      - 4 referencias para cada concepto
      - Años 2021–2025
      - Preferencia por DOI
    """
    resultados = {}

    # Teoría 1
    prompt_t1 = (
        "Imagina que para la variable teórica siguiente quieres sustentar una TEORÍA PRINCIPAL "
        "en el marco teórico de una revisión bibliográfica:\n\n"
        f"Variable: {variable1}\n\n"
        "Genera 3 referencias de artículos científicos en formato APA 7, en español o inglés, "
        "publicados entre 2021 y 2025, que sean apropiados para fundamentar una teoría relacionada con ese concepto.\n"
        "Condiciones:\n"
        "- Deben ser artículos de revistas científicas (no tesis, no capítulos de libro).\n"
        "- Incluye DOI o URL del artículo cuando sea posible.\n"
        "- Devuelve solo la lista de referencias, una por línea, sin explicaciones adicionales."
    )
    refs_t1_raw = gpt(prompt_t1)
    resultados["teoria1"] = [r.strip() for r in refs_t1_raw.split("\n") if r.strip()]

    # Teoría 2
    prompt_t2 = (
        "Imagina que para la siguiente variable aplicada o contextual quieres sustentar una SEGUNDA TEORÍA "
        "en el marco teórico de una revisión bibliográfica:\n\n"
        f"Variable: {variable2}\n\n"
        "Genera 3 referencias de artículos científicos en formato APA 7, en español o inglés, "
        "publicados entre 2021 y 2025, que sean apropiados para esa teoría.\n"
        "Condiciones similares:\n"
        "- Artículos de revistas científicas.\n"
        "- Incluye DOI o URL.\n"
        "- Solo la lista, una referencia por línea."
    )
    refs_t2_raw = gpt(prompt_t2)
    resultados["teoria2"] = [r.strip() for r in refs_t2_raw.split("\n") if r.strip()]

    # Variable 1 (concepto)
    prompt_v1 = (
        "Genera 4 referencias de artículos científicos en formato APA 7, en español o inglés, "
        "sobre el CONCEPTO teórico siguiente, pensadas para una revisión bibliográfica:\n\n"
        f"Concepto / variable 1: {variable1}\n\n"
        "Condiciones:\n"
        "- Artículos científicos de 2021 a 2025.\n"
        "- Incluye DOI o URL.\n"
        "- Solo la lista de referencias, una por línea."
    )
    refs_v1_raw = gpt(prompt_v1)
    resultados["variable1"] = [r.strip() for r in refs_v1_raw.split("\n") if r.strip()]

    # Variable 2 (concepto)
    prompt_v2 = (
        "Genera 4 referencias de artículos científicos en formato APA 7, en español o inglés, "
        "sobre el CONCEPTO aplicado o contextual siguiente, pensadas para una revisión bibliográfica:\n\n"
        f"Concepto / variable 2: {variable2}\n\n"
        "Condiciones:\n"
        "- Artículos científicos de 2021 a 2025.\n"
        "- Incluye DOI o URL.\n"
        "- Solo la lista de referencias, una por línea."
    )
    refs_v2_raw = gpt(prompt_v2)
    resultados["variable2"] = [r.strip() for r in refs_v2_raw.split("\n") if r.strip()]

    return resultados


def generate_reference_doc_phase2(
    titulo: str,
    pais: str,
    variable1: str,
    variable2: str,
    indexacion: str,
) -> str:
    """
    Genera el Word de PLANTILLA DE REFERENCIAS según la Segunda Parte:
      - Problemática mundial
      - Problemática latinoamericana (LATAM)
      - Problemática nacional – país
      - Teoría 1
      - Teoría 2
      - Variable 1
      - Variable 2
    Usando GPT para construir referencias APA 7.
    """
    doc = Document()

    # Encabezado principal
    doc.add_heading("Plantilla de referencias", level=0)
    titulo_resumido = resumir_titulo(titulo)
    doc.add_paragraph(titulo_resumido, style="Title")
    doc.add_paragraph(f"Indexación objetivo: {indexacion}")
    doc.add_paragraph("")

    # 1) Referencias institucionales (problemática)
    refs_problematica = generar_refs_institucionales(titulo, indexacion, pais)

    doc.add_heading("Problemática mundial", level=2)
    for ref in refs_problematica.get("mundial", []):
        doc.add_paragraph(ref, style="Normal")

    doc.add_paragraph("")

    doc.add_heading("Problemática latinoamericana (LATAM)", level=2)
    for ref in refs_problematica.get("latam", []):
        doc.add_paragraph(ref, style="Normal")

    doc.add_paragraph("")

    doc.add_heading(f"Problemática nacional – {pais}", level=2)
    for ref in refs_problematica.get("nacional", []):
        doc.add_paragraph(ref, style="Normal")

    doc.add_paragraph("")

    # 2) Referencias teóricas y conceptuales
    refs_teoricas = generar_refs_teoricas_y_conceptuales(variable1, variable2)

    doc.add_heading("Teoría 1", level=2)
    for ref in refs_teoricas.get("teoria1", []):
        doc.add_paragraph(ref, style="Normal")

    doc.add_paragraph("")

    doc.add_heading("Teoría 2", level=2)
    for ref in refs_teoricas.get("teoria2", []):
        doc.add_paragraph(ref, style="Normal")

    doc.add_paragraph("")

    doc.add_heading(f"Variable 1 – {variable1}", level=2)
    for ref in refs_teoricas.get("variable1", []):
        doc.add_paragraph(ref, style="Normal")

    doc.add_paragraph("")

    doc.add_heading(f"Variable 2 – {variable2}", level=2)
    for ref in refs_teoricas.get("variable2", []):
        doc.add_paragraph(ref, style="Normal")

   # Guardar documento con timestamp dentro de ./output
    output_dir = "./output"
    os.makedirs(output_dir, exist_ok=True)

    filename = f"referencias_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)

    doc.save(filepath)
    return filepath
