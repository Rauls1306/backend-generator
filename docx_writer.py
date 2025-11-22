from docx import Document


def save_article_to_docx(titulo, texto, ruta_archivo):
    """
    Crea un Word con el título arriba y luego el texto continuo.
    No agrega encabezados ni subtítulos adicionales.
    """
    try:
        doc = Document()

        if titulo:
            doc.add_heading(titulo, 0)

        for linea in texto.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())

        doc.save(ruta_archivo)
    except Exception as e:
        print("❌ Error al guardar el DOCX:", str(e))
        raise e
