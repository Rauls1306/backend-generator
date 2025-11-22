import os
import re
import uuid
from typing import Dict, List

import openai
import requests
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from discusion_conclusion import generar_discusion_y_conclusion
from docx_writer import save_article_to_docx
from generator import (extraer_variables_desde_titulo, generar_titulo_es,
                       generate_article, traducir_variable_al_ingles)
from metodologia_prisma import generar_metodologia_y_figura_prisma
from prisma_processor import (generar_excel_prisma, generar_word_prisma,
                              procesar_pdfs_por_fuente)
from reference_writer import generate_reference_doc_phase2
from resumen_final import (generar_resumen_y_pulido,
                           insertar_resumen_y_pulido_en_doc)
from search_articles_serpapi import (buscar_latindex, buscar_scielo,
                                     buscar_scopus,
                                     generar_doc_busqueda_latindex,
                                     generar_doc_busqueda_scielo,
                                     generar_doc_busqueda_scopus)

# Cargar .env y configurar API KEY
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
print("DEBUG OPENAI_API_KEY:", bool(openai.api_key))

app = FastAPI()

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # ajústalo en producción
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class GeneradorInput(BaseModel):
    nombre: str
    pais: str
    tema: str
    indexacion: str   # tipo de artículo / nivel de indexación


class ArticuloPDF(BaseModel):
    pdf_path: str      # ruta local al PDF (ej: "pdfs/scielo/art1.pdf")
    base_datos: str    # "Dialnet", "SciELO" o "Scopus"

class PrismaStatsItem(BaseModel):
    base_datos: str           # "Scopus", "SciELO", "Dialnet", etc.
    identificados: int        # número de registros identificados
    incluidos: int            # artículos finalmente incluidos
    excluidos: int            # registros/artículos excluidos


class MetodologiaPrismaRequest(BaseModel):
    ruta_articulo: str        # ej: "./output/articulo_base_....docx"
    tema: str
    pais: str
    indexacion: str           # ej: "Scopus Q3", "SciELO", "LatIndex"
    stats: List[PrismaStatsItem]

class PrismaRequest(BaseModel):
    indexacion: str           # "LatIndex", "SciELO", "Scopus Q3", etc.
    articulos: List[ArticuloPDF]

# Para el pipeline completo: solo necesitamos los datos base
class PipelineRequest(BaseModel):
    datos: GeneradorInput

class DiscusionConclusionRequest(BaseModel):
    ruta_articulo: str        # ej: "./output/articulo_base_....docx"
    tema: str
    pais: str
    indexacion: str           # ej: "Scopus Q3", "SciELO", "LatIndex"
    stats: List[PrismaStatsItem]

class ResumenFinalRequest(BaseModel):
    ruta_articulo: str
    tema: str
    pais: str
    indexacion: str

# ─────────────────────────────────────────────
# Helpers para pipeline
# ─────────────────────────────────────────────

def _slugify(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]+", "_", text)


def build_sciencedirect_pdf_url(url: str) -> str:
    """
    A partir de una URL de artículo de ScienceDirect, intenta construir la URL del PDF.
    - Convierte '/science/article/abs/pii/' en '/science/article/pii/'
    - Añade '/pdf' al final, sin querystring.
    """
    base = url.split("?", 1)[0].rstrip("/")

    # Normalizar /abs/pii/ a /pii/
    base = base.replace("/science/article/abs/pii/", "/science/article/pii/")

    # Si ya termina en /pdf, la dejamos así
    if base.endswith("/pdf"):
        return base

    return base + "/pdf"

def descargar_pdfs_desde_urls(
    urls: List[str],
    base_datos: str,
    indexacion: str,
    max_pdfs: int = 30,
) -> List[str]:
    """
    Descarga PDFs desde URLs (mejor esfuerzo).
    - Para ScienceDirect, intenta convertir la URL HTML en URL de PDF.
    - Solo guarda si el Content-Type contiene 'pdf' o si forzamos PDF para ScienceDirect.
    - Guarda en ./pdfs/<indexacion>/<base_datos>/
    - Devuelve la lista de rutas locales de los PDFs descargados.
    """
    pdf_dir = os.path.join("pdfs", _slugify(indexacion), _slugify(base_datos))
    os.makedirs(pdf_dir, exist_ok=True)
    pdf_paths: List[str] = []

    for i, url in enumerate(urls, start=1):
        if i > max_pdfs:
            break

        try:
            original_url = url

            # Si es ScienceDirect, construir URL del PDF
            if "sciencedirect.com" in url:
                url = build_sciencedirect_pdf_url(url)
                print(f"🔁 Transformando URL ScienceDirect a PDF:\n   {original_url}\n   -> {url}")

            print(f"⬇️ Descargando ({base_datos}) {url}")
            resp = requests.get(url, timeout=30)
            ctype = resp.headers.get("Content-Type", "").lower()

            # Regla de aceptación:
            # - Si content-type incluye 'pdf' => lo guardamos
            # - Si es ScienceDirect y status 200 => lo intentamos guardar igual (muchas veces es PDF)
            es_sciencedirect = "sciencedirect.com" in url

            if ("pdf" not in ctype) and not es_sciencedirect:
                print(f"⚠️ No parece PDF (Content-Type={ctype}), se omite: {url}")
                continue

            filename = f"{_slugify(base_datos)}_articulo_{i}.pdf"
            fpath = os.path.join(pdf_dir, filename)

            with open(fpath, "wb") as f:
                f.write(resp.content)

            pdf_paths.append(fpath)
            print(f"✅ PDF guardado: {fpath}")

        except Exception as e:
            print(f"❌ Error descargando {url}: {e}")

    print(f"✅ PDFs descargados para {base_datos}: {len(pdf_paths)}")
    return pdf_paths

# ─────────────────────────────────────────────
# 🟢 ETAPA 1: Generar artículo base
# ─────────────────────────────────────────────

@app.post("/generar-articulo")
def generar_articulo_endpoint(data: GeneradorInput):
    try:
        print("✅ Datos recibidos en /generar-articulo:")
        print(data.model_dump())

        resultado = generate_article(
            tema=data.tema,
            nivel=data.indexacion,
            pais=data.pais,
        )

        titulo = resultado["titulo"]
        texto_articulo = resultado["texto_articulo"]

        output_dir = "./output"
        os.makedirs(output_dir, exist_ok=True)

        filename = f"articulo_base_{uuid.uuid4().hex}.docx"
        ruta_archivo = os.path.join(output_dir, filename)

        save_article_to_docx(titulo, texto_articulo, ruta_archivo)

        return {
            "titulo": resultado["titulo"],
            "variable_1": resultado["variable_1"],
            "variable_2": resultado["variable_2"],
            "texto_articulo": resultado["texto_articulo"],
            "archivo_generado": ruta_archivo,
        }

    except Exception as e:
        print("❌ Error durante la generación del artículo:")
        print(str(e))
        raise HTTPException(status_code=500, detail=str(e))


# ─────────────────────────────────────────────
# 🟣 ETAPA 2: PLANTILLA DE REFERENCIAS
# ─────────────────────────────────────────────

@app.post("/generar-referencias")
def generar_referencias_endpoint(data: GeneradorInput):
    try:
        print("✅ Datos recibidos en /generar-referencias:")
        print(data.model_dump())

        resultado = generate_article(
            tema=data.tema,
            nivel=data.indexacion,
            pais=data.pais,
        )

        titulo = resultado["titulo"]
        variable_1 = resultado["variable_1"]
        variable_2 = resultado["variable_2"]

        ref_filename = generate_reference_doc_phase2(
            titulo=titulo,
            pais=data.pais,
            variable1=variable_1,
            variable2=variable_2,
            indexacion=data.indexacion,
        )

        return {
            "titulo": titulo,
            "variable_1": variable_1,
            "variable_2": variable_2,
            "archivo_referencias": ref_filename,
        }

    except Exception as e:
        print("❌ Error durante la generación de referencias:")
        print(str(e))
        raise HTTPException(status_code=500, detail=str(e))


# ─────────────────────────────────────────────
# 🟡 ETAPA 3: Búsqueda de artículos
# ─────────────────────────────────────────────

@app.post("/buscar-articulos")
def buscar_articulos_endpoint(data: GeneradorInput):
    try:
        print("✅ Datos recibidos en /buscar-articulos:")
        print(data.model_dump())

        index_lower = data.indexacion.lower()

        # 1) Generar título ES y variable 1 específica desde el título
        titulo_es = generar_titulo_es(data.tema, data.pais, data.indexacion)
        variable_1, _ = extraer_variables_desde_titulo(titulo_es)
        print("📌 Para búsqueda, variable_1:", variable_1)

        # 2) Según indexación, crear el DOCX de búsqueda
        if "latindex" in index_lower:
            doc_filename = generar_doc_busqueda_latindex(titulo_es, variable_1)

        elif "sci" in index_lower:  # SciELO
            doc_filename = generar_doc_busqueda_scielo(titulo_es, variable_1)

        elif "scopus" in index_lower:
            # Traducir la variable 1 al inglés para ScienceDirect
            variable_1_en = traducir_variable_al_ingles(variable_1)
            print("📌 variable_1 en inglés:", variable_1_en)
            doc_filename = generar_doc_busqueda_scopus(titulo_es, variable_1_en)

        else:
            raise HTTPException(
                status_code=400,
                detail=f"Indexación no soportada para búsqueda: {data.indexacion}",
            )

        return {
            "titulo": titulo_es,
            "variable_1": variable_1,
            "archivo_busqueda": doc_filename,
        }

    except HTTPException:
        raise
    except Exception as e:
        print("❌ Error durante la búsqueda de artículos:")
        print(str(e))
        raise HTTPException(status_code=500, detail=str(e))


# ─────────────────────────────────────────────
# 🔵 ETAPA 4: PRISMA (con PDFs manuales)
# ─────────────────────────────────────────────

@app.post("/generar-prisma")
def generar_prisma(req: PrismaRequest):
    try:
        # Agrupar PDFs por base de datos (Dialnet, SciELO, Scopus)
        grupos: Dict[str, List[str]] = {}
        for art in req.articulos:
            grupos.setdefault(art.base_datos, []).append(art.pdf_path)

        resultados = []

        for base, paths in grupos.items():
            print(f"🔎 Procesando {len(paths)} PDFs para base {base}...")
            articulos_info = procesar_pdfs_por_fuente(paths)
            word_path = generar_word_prisma(req.indexacion, base, articulos_info)
            excel_path = generar_excel_prisma(req.indexacion, base, articulos_info)

            resultados.append(
                {
                    "base_datos": base,
                    "num_articulos": len(articulos_info),
                    "word_prisma": word_path,
                    "excel_prisma": excel_path,
                }
            )

        return {
            "indexacion": req.indexacion,
            "resultados": resultados,
        }

    except Exception as e:
        print("❌ Error en /generar-prisma:", e)
        raise HTTPException(status_code=500, detail=str(e))


# ─────────────────────────────────────────────
# 🧩 ETAPA 5: METODOLOGÍA + FIGURA PRISMA (endpoint propio)
# ─────────────────────────────────────────────

@app.post("/generar-metodologia-prisma")
def generar_metodologia_prisma_endpoint(req: MetodologiaPrismaRequest):
    """
    Genera la sección 'Metodología' y el texto de la Figura PRISMA
    a partir de:
      - El artículo base (ruta_articulo)
      - Tema, país, nivel de indexación
      - Estadísticas PRISMA por base de datos (identificados, incluidos, excluidos)

    Inserta la metodología en el DOCX del artículo (al final),
    y devuelve los textos generados.
    """
    try:
        # Convertir lista de stats en diccionario como espera metodologia_prisma.py
        prisma_stats_dict: Dict[str, Dict[str, int]] = {}
        for item in req.stats:
            prisma_stats_dict[item.base_datos] = {
                "identificados": item.identificados,
                "incluidos": item.incluidos,
                "excluidos": item.excluidos,
            }

        print("✅ Datos recibidos en /generar-metodologia-prisma:")
        print("ruta_articulo:", req.ruta_articulo)
        print("tema:", req.tema)
        print("pais:", req.pais)
        print("indexacion:", req.indexacion)
        print("prisma_stats:", prisma_stats_dict)

        resultado = generar_metodologia_y_figura_prisma(
            ruta_articulo=req.ruta_articulo,
            tema=req.tema,
            pais=req.pais,
            indexacion=req.indexacion,
            prisma_stats=prisma_stats_dict,
        )

        return {
            "ruta_articulo_actualizado": req.ruta_articulo,
            "metodologia": resultado.get("metodologia", ""),
            "figura_prisma": resultado.get("figura_prisma", ""),
            "prisma_stats": prisma_stats_dict,
        }

    except Exception as e:
        print("❌ Error en /generar-metodologia-prisma:", e)
        raise HTTPException(status_code=500, detail=str(e))


# ─────────────────────────────────────────────
# 🧩 ETAPA 6: DISCUSIÓN + CONCLUSIONES (endpoint propio)
# ─────────────────────────────────────────────

@app.post("/generar-discusion-conclusion")
def generar_discusion_conclusion_endpoint(req: DiscusionConclusionRequest):
    """
    Genera las secciones 'Discusión' y 'Conclusiones' del artículo
    a partir de:
      - El artículo base (ruta_articulo)
      - Tema, país, nivel de indexación
      - Estadísticas PRISMA por base de datos (identificados, incluidos, excluidos)

    Inserta ambas secciones en el DOCX del artículo (al final),
    y devuelve los textos generados.
    """
    try:
        prisma_stats_dict: Dict[str, Dict[str, int]] = {}
        for item in req.stats:
            prisma_stats_dict[item.base_datos] = {
                "identificados": item.identificados,
                "incluidos": item.incluidos,
                "excluidos": item.excluidos,
            }

        print("✅ Datos recibidos en /generar-discusion-conclusion:")
        print("ruta_articulo:", req.ruta_articulo)
        print("tema:", req.tema)
        print("pais:", req.pais)
        print("indexacion:", req.indexacion)
        print("prisma_stats:", prisma_stats_dict)

        resultado = generar_discusion_y_conclusion(
            ruta_articulo=req.ruta_articulo,
            tema=req.tema,
            pais=req.pais,
            indexacion=req.indexacion,
            prisma_stats=prisma_stats_dict,
        )

        return {
            "ruta_articulo_actualizado": req.ruta_articulo,
            "discusion": resultado.get("discusion", ""),
            "conclusion": resultado.get("conclusion", ""),
            "prisma_stats": prisma_stats_dict,
        }

    except Exception as e:
        print("❌ Error en /generar-discusion-conclusion:", e)
        raise HTTPException(status_code=500, detail=str(e))

# ─────────────────────────────────────────────
# 🧩 PARTE 7 — RESUMEN FINAL + PULIDO
# ─────────────────────────────────────────────

@app.post("/generar-resumen-final")
def generar_resumen_final_endpoint(req: ResumenFinalRequest):
    try:
        print("✅ Generando resumen final y pulido...")

        resumen_es, resumen_en, texto_pulido = generar_resumen_y_pulido(
            ruta_articulo=req.ruta_articulo,
            tema=req.tema,
            pais=req.pais,
            indexacion=req.indexacion
        )

        ruta_final = insertar_resumen_y_pulido_en_doc(
            ruta_articulo=req.ruta_articulo,
            resumen_es=resumen_es,
            resumen_en=resumen_en,
            texto_pulido=texto_pulido
        )

        return {
            "ruta_articulo_final": ruta_final,
            "resumen_es": resumen_es,
            "resumen_en": resumen_en,
            "pulido": texto_pulido
        }

    except Exception as e:
        print("❌ Error en /generar-resumen-final:", e)
        raise HTTPException(status_code=500, detail=str(e))

# ─────────────────────────────────────────────
# 🚀 PIPELINE COMPLETO: TODO AUTOMÁTICO
# ─────────────────────────────────────────────

@app.post("/pipeline-completo")
def pipeline_completo(req: PipelineRequest):
    """
    Pipeline completo corregido:
    - PRISMA, METODOLOGÍA, DISCUSIÓN Y CONCLUSIÓN ahora se ejecutan SIEMPRE para Scopus.
    - El RESUMEN/ABSTRACT se reemplazan correctamente sin destruir encabezados.
    - El documento integrado ahora inserta PRISMA + METODOLOGÍA + DISCUSIÓN + REFERENCIAS dentro del DOC.
    """

    # ============================================================
    # 🔧 FIX 0 – Helper corregido: insertar capítulos en orden real
    # ============================================================
    def insert_section(doc: Document, title: str, content: str):
        """Inserta una sección (título + párrafo) al final del documento."""
        if not content:
            return

        doc.add_page_break()
        h = doc.add_heading(level=1)
        h.add_run(title)
        for p in content.split("\n"):
            if p.strip():
                doc.add_paragraph(p.strip())

    # ============================================================
    # 🔧 FIX 1 – Nuevo método de integración TOTAL
    # ============================================================
    def build_integrated_doc(
        ruta_base_final,
        metodologia_texto,
        figura_prisma_texto,
        discusion_texto,
        conclusion_texto,
        prisma_word_paths,
        ruta_referencias_doc,
        output_path,
    ) -> str:

        doc = Document(ruta_base_final)

        # → METODOLOGÍA
        insert_section(doc, "METODOLOGÍA", metodologia_texto)

        # → FIGURA PRISMA
        insert_section(doc, "Figura 1. Diagrama de flujo PRISMA", figura_prisma_texto)

        # → DISCUSIÓN
        insert_section(doc, "DISCUSIÓN", discusion_texto)

        # → CONCLUSIÓN
        insert_section(doc, "CONCLUSIÓN", conclusion_texto)

        # → INSERTAR WORDS DE PRISMA
        for path in prisma_word_paths:
            doc.add_page_break()
            try:
                extra = Document(path)
                for elem in extra.element.body:
                    doc.element.body.append(elem)
            except:
                print(f"⚠️ No se pudo anexar {path}")

        # → REFERENCIAS
        if ruta_referencias_doc and os.path.exists(ruta_referencias_doc):
            doc.add_page_break()
            h = doc.add_heading(level=1)
            h.add_run("REFERENCIAS")

            ref_doc = Document(ruta_referencias_doc)
            for elem in ref_doc.element.body:
                doc.element.body.append(elem)

        doc.save(output_path)
        return output_path

    # ============================================================
    # 🔧 FIX 2 – Postprocesador NO destruye Resumen/Abstract
    # ============================================================
    def replace_summary_abstract(ruta_articulo, resumen_es, resumen_en):
        doc = Document(ruta_articulo)

        # Encontrar secciones por encabezados
        def replace_block(header, new_text):
            found = False
            to_delete = []

            for i, p in enumerate(doc.paragraphs):
                if p.text.strip().upper() == header:
                    found = True
                    start = i
                    continue

                if found:
                    if p.text.strip().isupper() and len(p.text.strip()) < 25:
                        # siguiente encabezado → detener
                        break
                    to_delete.append(p)

            # Borrar contenido viejo
            for p in to_delete:
                p._element.getparent().remove(p._element)

            # Insertar nuevo texto
            anchor = doc.paragraphs[start]
            for line in new_text.split("\n"):
                new_p = anchor.insert_paragraph_before(line)
            return

        replace_block("RESUMEN", resumen_es)
        replace_block("ABSTRACT", resumen_en)

        new_path = ruta_articulo.replace(".docx", "_final.docx")
        doc.save(new_path)
        return new_path

    # ======================================================================
    # ======================================================================
    # 🔥🔥🔥  AQUÍ EMPIEZA EL PIPELINE ORIGINAL PERO CON ARREGLOS  🔥🔥🔥
    # ======================================================================
    # ======================================================================

    try:
        datos = req.datos
        index_lower = datos.indexacion.lower()

        output_dir = "./output"
        os.makedirs(output_dir, exist_ok=True)

        # ------------------------------------------------------------------
        # PARTE 1: Artículo base
        # ------------------------------------------------------------------
        resultado = generate_article(
            tema=datos.tema,
            nivel=datos.indexacion,
            pais=datos.pais,
        )

        titulo_es = resultado["titulo"]
        variable_1 = resultado["variable_1"]
        variable_2 = resultado["variable_2"]
        texto_articulo = resultado["texto_articulo"]

        ruta_articulo = os.path.join(
            output_dir, f"articulo_base_{uuid.uuid4().hex}.docx"
        )
        save_article_to_docx(titulo_es, texto_articulo, ruta_articulo)

        # ------------------------------------------------------------------
        # PARTE 2: Referencias
        # ------------------------------------------------------------------
        ruta_referencias_doc = generate_reference_doc_phase2(
            titulo=titulo_es,
            pais=datos.pais,
            variable1=variable_1,
            variable2=variable_2,
            indexacion=datos.indexacion,
        )

        # ------------------------------------------------------------------
        # PARTE 3: Búsqueda Scopus
        # ------------------------------------------------------------------
        bases_urls = {}
        variable_1_en = traducir_variable_al_ingles(variable_1)
        urls_scopus = buscar_scopus(variable_1_en)
        bases_urls["Scopus"] = urls_scopus
        busqueda_filename = generar_doc_busqueda_scopus(titulo_es, variable_1_en)

        # ------------------------------------------------------------------
        # PARTE 4: PRISMA
        # ------------------------------------------------------------------
        prisma_resultados = []
        prisma_stats = {}
        prisma_word_paths = []

        for base, urls in bases_urls.items():
            pdf_paths = descargar_pdfs_desde_urls(
                urls=urls,
                base_datos=base,
                indexacion=datos.indexacion,
            )

            articulos_info = procesar_pdfs_por_fuente(pdf_paths)

            prisma_stats[base] = {
                "identificados": len(urls),
                "incluidos": len(articulos_info),
                "excluidos": max(0, len(urls) - len(articulos_info)),
            }

            word_path = generar_word_prisma(datos.indexacion, base, articulos_info)
            excel_path = generar_excel_prisma(datos.indexacion, base, articulos_info)
            prisma_word_paths.append(word_path)

            prisma_resultados.append(
                {
                    "base_datos": base,
                    "num_articulos": len(articulos_info),
                    "word_prisma": word_path,
                    "excel_prisma": excel_path,
                }
            )

        # ------------------------------------------------------------------
        # PARTE 5: Metodología + Figura PRISMA
        # ------------------------------------------------------------------
        res_met = generar_metodologia_y_figura_prisma(
            ruta_articulo=ruta_articulo,
            tema=datos.tema,
            pais=datos.pais,
            indexacion=datos.indexacion,
            prisma_stats=prisma_stats,
        )
        metodologia_texto = res_met.get("metodologia", "")
        figura_prisma_texto = res_met.get("figura_prisma", "")

        # ------------------------------------------------------------------
        # PARTE 6: Discusión + Conclusión
        # ------------------------------------------------------------------
        res_dc = generar_discusion_y_conclusion(
            ruta_articulo=ruta_articulo,
            tema=datos.tema,
            pais=datos.pais,
            indexacion=datos.indexacion,
            prisma_stats=prisma_stats,
        )
        discusion_texto = res_dc.get("discusion", "")
        conclusion_texto = res_dc.get("conclusion", "")

        # ------------------------------------------------------------------
        # PARTE 7: Resumen + Polishing
        # ------------------------------------------------------------------
        resumen_es, resumen_en, texto_pulido = generar_resumen_y_pulido(
            ruta_articulo=ruta_articulo,
            tema=datos.tema,
            pais=datos.pais,
            indexacion=datos.indexacion,
        )

        ruta_articulo_final = replace_summary_abstract(
            ruta_articulo, resumen_es, resumen_en
        )

        # ------------------------------------------------------------------
        # INTEGRACIÓN REAL COMPLETA 🔥 FIX
        # ------------------------------------------------------------------
        ruta_doc_integrado = os.path.join(
            output_dir,
            f"articulo_integrado_{uuid.uuid4().hex}.docx"
        )

        ruta_doc_integrado = build_integrated_doc(
            ruta_base_final=ruta_articulo_final,
            metodologia_texto=metodologia_texto,
            figura_prisma_texto=figura_prisma_texto,
            discusion_texto=discusion_texto,
            conclusion_texto=conclusion_texto,
            prisma_word_paths=prisma_word_paths,
            ruta_referencias_doc=ruta_referencias_doc,
            output_path=ruta_doc_integrado,
        )

        # ------------------------------------------------------------------
        # RESPUESTA
        # ------------------------------------------------------------------
        return {
            "datos_entrada": datos.model_dump(),
            "articulo_base": {
                "titulo": titulo_es,
                "variable_1": variable_1,
                "variable_2": variable_2,
                "archivo": ruta_articulo,
            },
            "referencias": {"archivo_referencias": ruta_referencias_doc},
            "busqueda_articulos": {"archivo_busqueda": busqueda_filename},
            "prisma": {
                "ejecutado": True,
                "resultados": prisma_resultados,
                "stats": prisma_stats,
            },
            "metodologia": {"generada": True, "texto": metodologia_texto},
            "figura_prisma": {"texto": figura_prisma_texto},
            "discusion_conclusion": {
                "generada": True,
                "discusion": discusion_texto,
                "conclusion": conclusion_texto,
            },
            "resumen_final": {
                "generado": True,
                "ruta_articulo_final": ruta_articulo_final,
                "resumen_es": resumen_es,
                "resumen_en": resumen_en,
            },
            "documento_integrado": {
                "generado": True,
                "ruta_doc_integrado": ruta_doc_integrado,
            },
        }

    except Exception as e:
        print("❌ Error en pipeline:", e)
        raise HTTPException(status_code=500, detail=str(e))
