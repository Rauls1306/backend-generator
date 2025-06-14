import openai
from docx import Document
from datetime import datetime
import os

openai.api_key = os.getenv("OPENAI_API_KEY")

def gpt(prompt):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Eres un redactor académico SCOPUS Q1. Responde solo con el texto pedido, sin subtítulos ni encabezados. Usa estilo impersonal, académico, fluido."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.65,
            max_tokens=1800
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error al generar contenido: {str(e)}"

def generate_article(tema, nivel):
    doc = Document()
    doc.add_heading("Artículo generado automáticamente", 0)
    doc.add_paragraph(f"Tema: {tema}")
    doc.add_paragraph(f"Nivel de indexación: {nivel}")
    doc.add_paragraph("")

    titulo = gpt(f"Genera un título académico Scopus Q1, sin adornos, basado en el siguiente tema: {tema}")
    doc.add_heading(titulo, level=1)

    doc.add_heading("Introducción", level=2)

    # ORDEN REAL DE LA INTRODUCCIÓN

    # 1. Párrafo de contexto
    doc.add_paragraph(gpt(
        f"Redacta un párrafo tipo SCOPUS Q1 como el modelo: 'Hazme un párrafo como este…' contextualizando el tema: {tema}"
    ))

    # 2. Problema mundial
    doc.add_paragraph(gpt(
        f"Redacta un párrafo con datos cuantitativos reales y actualizados sobre el problema de '{tema}' a nivel mundial, sin citas ni mención de instituciones."
    ))

    # 3. Problema en América Latina
    doc.add_paragraph(gpt(
        f"Redacta un párrafo con datos cuantitativos reales y actuales sobre el problema de '{tema}' en América Latina. Estilo académico SCOPUS, sin fuentes visibles."
    ))

    # 4. Problema en Perú
    doc.add_paragraph(gpt(
        f"Redacta un párrafo sobre el problema de '{tema}' en Perú, con datos cuantitativos reales pero sin citar autores ni instituciones."
    ))

    # 5. Problema, causas y consecuencias
    doc.add_paragraph(gpt(
        f"Redacta un párrafo SCOPUS Q1 que explique el problema de '{tema}' con sus causas y consecuencias. Texto fluido, sin listas."
    ))

    # 6. Justificación
    doc.add_paragraph(gpt(
        f"Redacta un párrafo académico tipo SCOPUS que justifique la importancia del artículo sobre: {tema}. Estilo impersonal, convincente, sin adornos."
    ))

    # MARCO TEÓRICO
    doc.add_heading("Marco teórico", level=2)

    # Teoría 1
    doc.add_paragraph(gpt(
        "Redacta un texto de 200 palabras sobre la Teoría de la Calidad Educativa Universitaria, incluyendo autor principal, contexto histórico y explicación, en prosa."
    ))

    # Teoría 2
    doc.add_paragraph(gpt(
        "Redacta un texto de 200 palabras sobre la Teoría de la Docencia Universitaria, con autor destacado, época y aplicación académica. Solo texto en prosa."
    ))

    # Variable 1: competencias docentes universitarios
    doc.add_paragraph(gpt(
        "Redacta tres párrafos sobre la variable 'Competencias en docentes universitarios'. El primero con la definición, el segundo con características y el tercero con tipos."
    ))

    # Variable 2: educación superior de calidad
    doc.add_paragraph(gpt(
        "Redacta tres párrafos sobre la variable 'Educación superior de calidad'. El primero con la definición, el segundo con criterios y el tercero con enfoques actuales."
    ))

    filename = f"/tmp/articulo_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    doc.save(filename)
    return filename