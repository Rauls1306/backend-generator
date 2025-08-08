import openai
import time
import os
from docx import Document
from docx_writer import save_article_to_docx
from datetime import datetime
from reference_writer import generate_reference_doc
from citation_generator import CitationGenerator
from generator_utils import (
    generate_apa_references,
    generate_textual_citations,
    insert_citations_into_text
)


openai.api_key = os.environ.get("OPENAI_API_KEY")

def gpt(prompt):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.65,
            max_tokens=2000
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error al generar contenido: {str(e)}"

def extract_concepts(titulo):
    prompt = (
        f"Del siguiente título académico: {titulo}, extrae dos conceptos principales: uno técnico desde la profesión del usuario y otro contextual desde el entorno o sector involucrado. "
        f"Devuélvelos sin comillas, en minúsculas, sin numeración, separados por salto de línea."
    )
    resultado = gpt(prompt)
    return [v.strip() for v in resultado.split("\n") if v.strip()]

def generate_article(tema, nivel, pais):
    doc = Document()
    doc.add_heading("Artículo generado automáticamente", 0)
    doc.add_paragraph(f"Tema ingresado: {tema}")
    doc.add_paragraph(f"Nivel de indexación: {nivel}")
    doc.add_paragraph("")

    # TÍTULO
    prompt_titulo = (
        f"A partir del siguiente input informal: '{tema}', genera un título académico formal con redacción Scopus. "
        f"Debe contener una combinación entre un concepto técnico derivado de la carrera y otro del entorno. "
        f"No repitas frases del input, no uses comillas ni fórmulas genéricas como 'un estudio sobre' o 'intersección entre'."
    )
    titulo = gpt(prompt_titulo)
    doc.add_heading(titulo, level=1)
    time.sleep(2)
    
    # VARIABLES
    variables = extract_concepts(titulo)
    print("📌 Variables extraídas del título:", variables)
    variable_1 = variables[0] if len(variables) > 0 else "variable técnica no identificada"
    variable_2 = variables[1] if len(variables) > 1 else "variable contextual no identificada"

    doc.add_heading(titulo, level=1)
    time.sleep(2)

    # CONTEXTO
    contexto = gpt(
        f"Redacta un texto así sobre la problemática del artículo titulado '{titulo}', mismo tamaño, mismo 1 párrafo, recuerda, sin usar datos cuantitativos, NO MENCIONES EL TITULO DE LA INVESTIGACION. HAZLO COMO ESTE MODELO: Los polifenoles han demostrado tener un impacto positivo en la reducción de los niveles lipídicos..."
    )
    if "Error al generar contenido" in contexto:
        raise ValueError("❌ GPT no generó el contexto correctamente.")
    time.sleep(2)

    # MUNDIAL / LATAM / PAÍS
    mundial_latam_peru = gpt(
        f"Redacta un texto de 3 párrafos, c/u de 100 palabras, todo estilo scopus q1, sobre la problemática del artículo titulado '{titulo}'. "
        f"Cada párrafo es por un nivel: el primer párrafo nivel global o mundial, segundo nivel LATAM, tercero nivel nacional del país {pais}. "
        f"Cada párrafo debe tener 3 datos cuantitativos (solo uno porcentual, los otros 2 numericos, IMPORTANTEeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeeee). "
        f"No incluyas citas ni menciones a instituciones (IMPORTANTISIMO) ni ambigüedades como 'cerca de' o 'casi'. No uses conectores de cierre. "
        f"Cada párrafo debe iniciar mencionando el nivel (ejemplo: A nivel global, En Latinoamérica, En el contexto de {pais}). "
        f"Además, cada párrafo debe tener 2 datos cualitativos. TODA SOLO INFORMACION DE LOS ULTIMOS 5 AÑOS. importante, no uses la palabra \"CUALITATIVA\" ni similares"
    )
    print("🧪 Respuesta de GPT para niveles:", mundial_latam_peru)
    if "Error al generar contenido" in mundial_latam_peru:
        raise ValueError("❌ GPT no generó el niveles_raw correctamente.")
    niveles = [p.strip() for p in mundial_latam_peru.split("\n\n") if p.strip()]
    while len(niveles) < 3:
        niveles.append("⚠️ Contenido faltante generado automáticamente.")
    mundial, latam, peru = niveles[:3]
    time.sleep(2)

    # PROBLEMA
    problema = gpt(
        f"Redáctame un párrafo como este sobre problema, causas y consecuencias sobre la problemática del artículo titulado '{titulo}', en 90 palabras, redactado como para scopus q1, sin datos cuantitativos, sin citas, sin tanta puntuación o separación en las oraciones, que sea un párrafo fluido..."
    )
    if "Error al generar contenido" in problema:
        raise ValueError("❌ GPT no generó el problema correctamente.")
    time.sleep(2)

    # JUSTIFICACIÓN
    justificacion = gpt(
        f"Redacta un párrafo de justificación, por relevancia, importancia, etc., de 100 palabras, estilo scopus q1, que empiece con la frase 'se justifica'. No menciones el título del artículo en esta justificación."
    )
    if "Error al generar contenido" in justificacion:
        raise ValueError("❌ GPT no generó el justificacion correctamente.")
    time.sleep(2)

    # TEORÍAS
    teorias = gpt(
        f"Busca dos teorías para una investigación que combina '{variable_1}' y '{variable_2}'. "
        f"Redacta un párrafo de 150 palabras para cada una. Menciona nombre de la teoría, padre, explicación. Sin subtítulos ni conectores de cierre."
    )
    if "Error al generar contenido" in teorias:
        raise ValueError(f"❌ Error generado por GPT al obtener teorías: {teorias}")

    teorias_split = teorias.split("\n\n")
    while len(teorias_split) < 2:
        teorias_split.append("⚠️ Teoría faltante")
    teoria1, teoria2 = teorias_split[:2]
    time.sleep(2)

    # CONCEPTOS
    conceptos = gpt(
        f"A partir de los conceptos '{variable_1}' y '{variable_2}', redacta para cada uno dos párrafos de 100 palabras. "
        f"Cada párrafo debe iniciar con un conector de adición (como: de manera concordante, en consonancia con lo anterior...). "
        f"No menciones que son variables, no uses títulos ni encabezados. Todo en prosa académica continua."
    )
    if "Error al generar contenido" in conceptos:
        raise ValueError("❌ GPT no generó el conceptos correctamente.")
    conceptos_divididos = conceptos.split("\n\n")
    while len(conceptos_divididos) < 5:
        conceptos_divididos.append("")

    concepto1_p1, concepto1_p2 = conceptos_divididos[0], conceptos_divididos[1]
    concepto2_p1, concepto2_p2, concepto2_p3 = conceptos_divididos[2], conceptos_divididos[3], conceptos_divididos[4]

    # GENERAR ARTÍCULO COMPLETO
    generated_text = {
        "contexto": contexto,
        "mundial": mundial,
        "latam": latam,
        "peru": peru,
        "problema": problema,
        "justificacion": justificacion,
        "teoria1": teoria1,
        "teoria2": teoria2,
        "concepto1_p1": concepto1_p1,
        "concepto1_p2": concepto1_p2,
        "concepto2_p1": concepto2_p1,
        "concepto2_p2": concepto2_p2,
        "concepto2_p3": concepto2_p3,
        "variable1_title": variable_1,
        "variable2_title": variable_2,
    }

    citation_generator = CitationGenerator(title=titulo, generated_text=generated_text)
    citation_generator.generate_all_citations()
    cited_texts = citation_generator.insert_all_citations()
    reference_filename = generate_reference_doc(
    titulo=titulo,
    pais=pais,
    referencias_dict=cited_texts 
)

    # Concatenar el texto final con citas insertadas
    final_article = ""
    for key in [
        "contexto", "mundial", "latam", "peru", "problema", "justificacion",
        "teoria1", "teoria2",
        "concepto1_p1", "concepto1_p2",
        "concepto2_p1", "concepto2_p2", "concepto2_p3"
    ]:
        value = cited_texts.get(key)
        if isinstance(value, str):
            final_article += value.strip() + "\n\n"
        else:
            print(f"⚠️ Advertencia: El bloque '{key}' está ausente o no es texto válido. Valor: {value}")

    final_article = final_article.strip()
    if not final_article or len(final_article) < 100:
        raise ValueError("❌ El contenido del artículo es muy corto o vacío. No se generará el Word.")

    filename = f"articulo_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    save_article_to_docx(final_article, filename)

    # Generar documento de referencias usando el texto citado
    # Agrupar los textos relacionados a las variables
    variable1_texts = [
        cited_texts.get("concepto1_p1", ""),
        cited_texts.get("concepto1_p2", "")
    ]

    variable2_texts = [
        cited_texts.get("concepto2_p1", ""),
        cited_texts.get("concepto2_p2", ""),
        cited_texts.get("concepto2_p3", "")
    ]

    # Preparar diccionario para referencias
    referencias_dict = {
        "mundial": [cited_texts.get("mundial", "")],
        "latam": [cited_texts.get("latam", "")],
        "peru": [cited_texts.get("peru", "")],
        "teoria1": [cited_texts.get("teoria1", "")],
        "teoria2": [cited_texts.get("teoria2", "")],
        "variable1": [t for t in variable1_texts if t],
        "variable2": [t for t in variable2_texts if t],
    }

    reference_filename = generate_reference_doc(
        titulo=titulo,
        pais=pais,
        referencias_dict=referencias_dict
    )

    print(f"📎 Plantilla de referencias generada: {reference_filename}")
    return filename, reference_filename
