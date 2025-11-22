# search_articles_serpapi.py
import os
from datetime import datetime
from typing import Dict, List, Optional

import requests
from docx import Document
from dotenv import load_dotenv

# 👇 Cargar .env aquí, independiente de main.py
load_dotenv()


def _get_serpapi_key() -> Optional[str]:
    """
    Acepta nombres alternativos para maximizar compatibilidad con .env.
    Prioridad: SERPAPI_KEY > SERP_API_KEY
    """
    return os.getenv("SERPAPI_KEY") or os.getenv("SERP_API_KEY")


SERPAPI_KEY: Optional[str] = _get_serpapi_key()
print("DEBUG SERPAPI_KEY cargada (search_articles_serpapi):", bool(SERPAPI_KEY))  # True si existe, False si no


def _ensure_serpapi_key():
    """
    Lanza error claro si la API key no está configurada.
    """
    if not SERPAPI_KEY:
        raise RuntimeError(
            "❌ SERPAPI_KEY no está configurada. "
            "Agrega SERPAPI_KEY=tu_clave o SERP_API_KEY=tu_clave en el archivo .env "
            "en la raíz del proyecto."
        )


def buscar_en_google(query: str, dominio: str, num: int) -> List[str]:
    """
    Realiza búsqueda REAL con SerpAPI (Google) vía HTTP.
    Filtra solo resultados del dominio indicado.
    Devuelve URLs REALES y limpias.
    """
    _ensure_serpapi_key()

    params = {
        "engine": "google",
        "q": query,
        "num": num * 3,   # pedimos más para luego filtrar
        "api_key": SERPAPI_KEY,
    }

    resp = requests.get("https://serpapi.com/search", params=params, timeout=30)
    resp.raise_for_status()
    results = resp.json()

    urls: List[str] = []

    organic = results.get("organic_results", [])
    for r in organic:
        link = r.get("link", "")
        if dominio.lower() in link.lower():
            urls.append(link)
        if len(urls) >= num:
            break

    return urls


# ─────────────────────────────────────────────
# LATINDEX (20 Dialnet)
# ─────────────────────────────────────────────

def buscar_latindex(variable1: str) -> List[str]:
    query = f'"{variable1}" site:dialnet.unirioja.es'
    urls = buscar_en_google(query, "dialnet", 20)
    return urls


# ─────────────────────────────────────────────
# SCIELO (15 SciELO + 15 Dialnet)
# ─────────────────────────────────────────────

def buscar_scielo(variable1: str) -> Dict[str, List[str]]:
    scielo = buscar_en_google(f'"{variable1}" site:scielo.org', "scielo", 15)
    dialnet = buscar_en_google(f'"{variable1}" site:dialnet.unirioja.es', "dialnet", 15)
    return {"scielo": scielo, "dialnet": dialnet}


# ─────────────────────────────────────────────
# SCOPUS Q3–Q4 (30 ScienceDirect)
# ─────────────────────────────────────────────

def buscar_scopus(variable1_en: str) -> List[str]:
    query = f'"{variable1_en}" site:sciencedirect.com'
    urls = buscar_en_google(query, "sciencedirect.com", 30)
    return urls


# ─────────────────────────────────────────────
# GENERACIÓN DE WORDS AUXILIARES
# ─────────────────────────────────────────────

def generar_doc_busqueda_latindex(titulo: str, variable1: str) -> str:
    urls = buscar_latindex(variable1)

    doc = Document()
    doc.add_heading("Artículos para revisión – LatIndex (Dialnet)", level=0)
    doc.add_paragraph(titulo, style="Title")
    doc.add_paragraph(f"Variable 1 específica: {variable1}\n")

    for u in urls:
        doc.add_paragraph(u)

    output_dir = "./output"
    os.makedirs(output_dir, exist_ok=True)

    filename = f"busqueda_latindex_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def generar_doc_busqueda_scielo(titulo: str, variable1: str) -> str:
    r = buscar_scielo(variable1)

    doc = Document()
    doc.add_heading("Artículos para revisión – SciELO", level=0)
    doc.add_paragraph(titulo, style="Title")
    doc.add_paragraph(f"Variable 1 específica: {variable1}\n")

    doc.add_heading("SciELO (15)", level=2)
    for u in r["scielo"]:
        doc.add_paragraph(u)

    doc.add_heading("Dialnet (15)", level=2)
    for u in r["dialnet"]:
        doc.add_paragraph(u)

    output_dir = "./output"
    os.makedirs(output_dir, exist_ok=True)

    filename = f"busqueda_scielo_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def generar_doc_busqueda_scopus(titulo: str, variable1_en: str) -> str:
    urls = buscar_scopus(variable1_en)

    doc = Document()
    doc.add_heading("Artículos para revisión – Scopus Q3–Q4 (ScienceDirect)", level=0)
    doc.add_paragraph(titulo, style="Title")
    doc.add_paragraph(f"Variable 1 específica (inglés): {variable1_en}\n")

    for u in urls:
        doc.add_paragraph(u)

    output_dir = "./output"
    os.makedirs(output_dir, exist_ok=True)

    filename = f"busqueda_scopus_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath