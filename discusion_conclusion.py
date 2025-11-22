# discusion_conclusion.py
from typing import Dict

import openai
from docx import Document


def _gpt(prompt: str, max_tokens: int = 1500, temperature: float = 0.35) -> str:
    """
    Wrapper simple para ChatCompletion (openai==0.28).
    Asume que openai.api_key ya está configurada en main.py.
    """
    resp = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    return resp.choices[0].message.content.strip()


def _build_prisma_summary_text(prisma_stats: Dict[str, Dict[str, int]]) -> str:
    """
    prisma_stats: {
      "Scopus": {"identificados": 30, "incluidos": 9, "excluidos": 21},
      "SciELO": {...},
      ...
    }
    """
    lines = []
    for base, stats in prisma_stats.items():
        ident = stats.get("identificados", 0)
        incl = stats.get("incluidos", 0)
        excl = stats.get("excluidos", 0)
        lines.append(
            f"- {base}: {ident} registros identificados, {excl} excluidos y {incl} artículos incluidos en la revisión."
        )
    return "\n".join(lines)


def _extract_article_snippet(ruta_articulo: str, max_chars: int = 8000) -> str:
    """
    Extrae el texto del artículo Word para dar contexto a GPT.
    Toma todos los párrafos y los recorta a max_chars.
    """
    try:
        doc = Document(ruta_articulo)
        textos = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                textos.append(t)
        full_text = "\n".join(textos)
        if len(full_text) > max_chars:
            return full_text[:max_chars]
        return full_text
    except Exception as e:
        print(f"⚠️ No se pudo extraer texto del DOCX ({ruta_articulo}): {e}")
        return ""


def generar_discusion_y_conclusion(
    ruta_articulo: str,
    tema: str,
    pais: str,
    indexacion: str,
    prisma_stats: Dict[str, Dict[str, int]],
) -> Dict[str, str]:
    """
    Genera:
    - Sección 'Discusión'
    - Sección 'Conclusiones'

    usando:
    - Tema, país, nivel de indexación
    - Resumen PRISMA (identificados, incluidos, excluidos)
    - Un fragmento del artículo (para mantener coherencia)

    Inserta ambas secciones en el DOCX del artículo (al final).
    """
    prisma_resumen_texto = _build_prisma_summary_text(prisma_stats)
    articulo_snippet = _extract_article_snippet(ruta_articulo)

    prompt = f"""
Eres un experto en redacción científica en español. A partir de la información del estudio, un fragmento del artículo y el resumen del flujo PRISMA, redacta:

1. La sección 'Discusión' de un artículo de revisión bibliográfica.
2. La sección 'Conclusiones' del mismo artículo.

DATOS DEL ESTUDIO
- Tipo de trabajo: revisión bibliográfica.
- Tema central: {tema}.
- País de enfoque principal: {pais}.
- Nivel de indexación: {indexacion}.

RESUMEN PRISMA POR BASE DE DATOS
{prisma_resumen_texto}

FRAGMENTO DEL ARTÍCULO (para contexto, no lo repitas literal):
\"\"\"
{articulo_snippet}
\"\"\"

INSTRUCCIONES PARA LA DISCUSIÓN
- Redactar entre 4 y 6 párrafos.
- Analizar críticamente los hallazgos más relevantes de los artículos incluidos.
- Comparar tendencias generales (por ejemplo: avances, brechas, limitaciones, implicancias para la práctica clínica o la gestión en {pais} o Latinoamérica).
- No inventar autores ni títulos específicos; habla de "los estudios revisados", "la literatura consultada", etc.
- Mencionar de forma general cómo el número y el tipo de artículos incluidos (según PRISMA) condicionan la solidez de las conclusiones.
- Mantener un tono académico, argumentativo y coherente.

INSTRUCCIONES PARA LAS CONCLUSIONES
- Redactar entre 2 y 3 párrafos.
- Resumir los aportes centrales de la revisión en relación con el tema '{tema}' en el contexto de {pais}.
- Incluir una frase sobre las implicancias prácticas o teóricas de los hallazgos.
- Incluir una frase sobre recomendaciones para futuras investigaciones o líneas de trabajo.
- No introducir información nueva que no se haya mencionado de forma general en la discusión.

FORMATO DE RESPUESTA (MUY IMPORTANTE)
Devuelve el resultado en este formato EXACTO:

DISCUSION:
[escribe aquí el texto completo de la sección Discusión]

CONCLUSION:
[escribe aquí el texto completo de la sección Conclusiones]

No agregues ningún otro encabezado ni comentarios fuera de este formato.
"""

    respuesta = _gpt(prompt)

    # Parseo simple
    discusion_texto = ""
    conclusion_texto = ""

    parts = respuesta.split("CONCLUSION:")
    if len(parts) == 2:
        parte_disc = parts[0]
        parte_conc = parts[1]

        if "DISCUSION:" in parte_disc:
            discusion_texto = parte_disc.split("DISCUSION:", 1)[1].strip()
        else:
            discusion_texto = parte_disc.strip()

        conclusion_texto = parte_conc.strip()
    else:
        # Fallback si el modelo no respeta el formato
        discusion_texto = respuesta.strip()
        conclusion_texto = ""

    # Insertar en el DOCX del artículo
    try:
        doc = Document(ruta_articulo)

        # Sección "RESULTADOS Y DISCUSIÓN"
        doc.add_heading("RESULTADOS Y DISCUSIÓN", level=1)
        for parrafo in discusion_texto.split("\n"):
            if parrafo.strip():
                doc.add_paragraph(parrafo.strip())

        # Sección "CONCLUSIONES"
        doc.add_heading("CONCLUSIONES", level=1)
        for parrafo in conclusion_texto.split("\n"):
            if parrafo.strip():
                doc.add_paragraph(parrafo.strip())

        doc.save(ruta_articulo)
    except Exception as e:
        print(f"❌ Error insertando Discusión/Conclusiones en el DOCX: {e}")

    return {
        "discusion": discusion_texto,
        "conclusion": conclusion_texto,
    }
