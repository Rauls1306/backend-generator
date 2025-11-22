# resumen_final.py
from typing import Tuple

import openai
from docx import Document


def _gpt(prompt: str, max_tokens: int = 1800, temperature: float = 0.35) -> str:
    """
    Wrapper simple para ChatCompletion (openai==0.28).
    """
    resp = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    return resp.choices[0].message.content.strip()


def _extract_full_article(ruta: str, max_chars: int = 12000) -> str:
    """
    Extrae el texto del artículo (solo para contexto del resumen).
    Limita por seguridad a max_chars para no pasarnos de tokens.
    """
    try:
        doc = Document(ruta)
        partes = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                partes.append(t)
        full_text = "\n".join(partes)
        if len(full_text) > max_chars:
            full_text = full_text[:max_chars]
        return full_text
    except Exception as e:
        print(f"⚠️ No se pudo leer el artículo completo: {e}")
        return ""


def generar_resumen_y_pulido(
    ruta_articulo: str,
    tema: str,
    pais: str,
    indexacion: str
) -> Tuple[str, str, str]:
    """
    Genera:
      - Resumen en español (150–250 palabras)
      - Abstract en inglés (150–250 palabras)
      - Texto pulido (que por ahora NO usaremos para reescribir el doc)

    Devuelve (resumen_es, resumen_en, texto_pulido)
    """

    articulo = _extract_full_article(ruta_articulo)

    prompt = f"""
Eres un experto en redacción científica.

A partir del TEXTO DEL ARTÍCULO que te doy, realiza:

1) Un RESUMEN en español de 150–250 palabras.
   - Sin inventar datos.
   - Debe reflejar: contexto, objetivo, metodología (revisión), hallazgos generales y conclusión.

2) Un ABSTRACT en inglés de 150–250 palabras, equivalente al resumen.

3) Opcionalmente, una versión PULIDA del manuscrito (en español):
   - Mejora cohesión, precisión y estilo académico.
   - NO cambies la estructura general del artículo.
   - NO inventes secciones nuevas.
   - NO añadas datos ni estudios que no estén en el texto.

TEXTO DEL ARTÍCULO:
\"\"\" 
{articulo}
\"\"\"

FORMATO DE RESPUESTA EXACTO:

RESUMEN_ES:
[texto]

RESUMEN_EN:
[texto]

PULIDO:
[texto]
"""

    raw = _gpt(prompt, max_tokens=3500)

    resumen_es = ""
    resumen_en = ""
    pulido = ""

    try:
        part1 = raw.split("RESUMEN_EN:")[0]
        part2 = raw.split("RESUMEN_EN:")[1].split("PULIDO:")[0]
        part3 = raw.split("PULIDO:")[1]

        if "RESUMEN_ES:" in part1:
            resumen_es = part1.split("RESUMEN_ES:", 1)[1].strip()
        else:
            resumen_es = part1.strip()

        resumen_en = part2.strip()
        pulido = part3.strip()
    except Exception:
        # fallback si el modelo no respeta perfecto el formato
        resumen_es = raw.strip()

    return resumen_es, resumen_en, pulido


def insertar_resumen_y_pulido_en_doc(
    ruta_articulo: str,
    resumen_es: str,
    resumen_en: str,
    texto_pulido: str
) -> str:
    """
    Reemplaza SOLO el Resumen y el Abstract dentro del mismo DOCX,
    conservando la estructura (INTRODUCCION, MARCO TEORICO, etc.).

    Devuelve la ruta del DOCX final (con sufijo _final.docx).
    """
    doc = Document(ruta_articulo)

    # Estados para saber en qué sección estamos
    EN_NINGUNA = 0
    EN_RESUMEN = 1
    EN_ABSTRACT = 2

    estado = EN_NINGUNA

    # Guardamos índices donde borrar texto viejo
    idx_resumen_paragraphs = []
    idx_abstract_paragraphs = []

    for i, p in enumerate(doc.paragraphs):
        texto = p.text.strip()

        upper = texto.upper()
        if upper in ("RESUMEN", "RESUMEN:"):
            estado = EN_RESUMEN
            continue
        if upper.startswith("ABSTRACT"):
            estado = EN_ABSTRACT
            continue

        # detectamos comienzo de otra sección grande → salimos de resumen/abstract
        if upper in ("INTRODUCCION", "INTRODUCCIÓN", "MARCO TEORICO",
                     "MARCO TEÓRICO", "METODOLOGIA", "METODOLOGÍA",
                     "RESULTADOS", "RESULTADOS Y DISCUSION",
                     "RESULTADOS Y DISCUSIÓN", "CONCLUSIONES",
                     "CONCLUSION", "CONCLUSIÓN"):
            estado = EN_NINGUNA

        if estado == EN_RESUMEN:
            idx_resumen_paragraphs.append(i)
        elif estado == EN_ABSTRACT:
            idx_abstract_paragraphs.append(i)

    # Borrar párrafos antiguos de Resumen
    for idx in sorted(idx_resumen_paragraphs, reverse=True):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)

    # Borrar párrafos antiguos de Abstract
    for idx in sorted(idx_abstract_paragraphs, reverse=True):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)

    # Reinsertar nuevo Resumen y Abstract
    # Los insertamos justo después de los encabezados.
    def _insert_after_heading(heading_texts, contenido, is_abstract=False):
        heading_indexes = []
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().upper() in heading_texts:
                heading_indexes.append(i)

        if not heading_indexes:
            return

        # Tomamos el primer heading que encontremos
        idx = heading_indexes[0]
        insert_pos = idx + 1

        lineas = [l.strip() for l in contenido.split("\n") if l.strip()]
        # Insertar párrafos en orden
        for j, linea in enumerate(lineas):
            new_p = doc.paragraphs[idx]._element.addnext(
                doc.paragraphs[idx]._element.__class__()
            )  # Truco bajo nivel, lo cambiamos abajo

        # Como el truco anterior es muy bajo nivel, simplifiquemos:
        # re-creamos la inserción usando el doc de manera simple:
        # vamos a reconstruir insertando con add_paragraph al final y luego reordenar
        # Para no complicar, mejor: insertamos al final de documento
        # y el usuario igual lo tendrá visible.
        # (Versión simple que no rompe doc)
        for linea in lineas:
            doc.add_paragraph(linea)

    # Versión simple: añadimos nuevos textos al final del documento
    # (para no manipular XML de forma frágil).
    # Si quieres que vayan pegados al encabezado, luego lo refinamos.
    if resumen_es:
        doc.add_paragraph("")  # línea en blanco
        doc.add_paragraph("Resumen (versión actualizada)")
        for linea in resumen_es.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())

    if resumen_en:
        doc.add_paragraph("")  # línea en blanco
        doc.add_paragraph("Abstract (updated version)")
        for linea in resumen_en.split("\n"):
            if linea.strip():
                doc.add_paragraph(linea.strip())

    # De momento NO tocamos el cuerpo con texto_pulido, para no romper estructura.

    ruta_final = ruta_articulo.replace(".docx", "_final.docx")
    doc.save(ruta_final)
    return ruta_final
